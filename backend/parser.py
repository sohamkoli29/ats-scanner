"""
parser.py - Handles text extraction from PDF, DOCX, and plain text files.
"""

import io
import re
import PyPDF2
import docx
from pdfminer.high_level import extract_text as pdfminer_extract


def extract_text_from_pdf(file_stream):
    """Extract text from a PDF file using PyPDF2 with pdfminer fallback."""
    text = ""
    try:
        reader = PyPDF2.PdfReader(file_stream)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception:
        pass

    # Fallback to pdfminer if PyPDF2 yields little text
    if len(text.strip()) < 100:
        try:
            file_stream.seek(0)
            text = pdfminer_extract(file_stream)
        except Exception:
            pass

    return text



def extract_text_from_docx(file_stream):
    """Extract text from a DOCX file including tables."""
    try:
        doc = docx.Document(file_stream)
        texts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text.strip())

        # Headers & Footers
        for section in doc.sections:
            for para in section.header.paragraphs + section.footer.paragraphs:
                if para.text.strip():
                    texts.append(para.text)

        return "\n".join(texts)
    except Exception as e:
        return ""


def extract_text_from_txt(file_stream):
    """Extract text from a plain text file."""
    try:
        raw = file_stream.read()
        return raw.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def parse_resume(file_stream, filename: str) -> str:
    """
    Detect file type and extract raw text.

    Args:
        file_stream: File-like object
        filename:    Original filename (used to detect extension)

    Returns:
        Extracted plain text string
    """
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        text = extract_text_from_pdf(file_stream)
    elif ext in ("docx", "doc"):
        text = extract_text_from_docx(file_stream)
    else:
        text = extract_text_from_txt(file_stream)

    return clean_text(text)


def clean_text(text: str) -> str:
    """Basic text normalisation."""
    # Collapse multiple whitespace / newlines
    text = re.sub(r"\s+", " ", text)
    # Remove non-ASCII except common punctuation
    text = re.sub(r"[^\x00-\x7F]+", " ", text)
    return text.strip()
